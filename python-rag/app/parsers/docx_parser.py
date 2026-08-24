import io
from typing import List
import docx
from app.parsers.base import BaseParser, ParsedDocument, ParsedSection


class DOCXParser(BaseParser):
    def parse(self, file_bytes: bytes, file_name: str) -> ParsedDocument:
        stream = io.BytesIO(file_bytes)
        doc = docx.Document(stream)
        
        sections: List[ParsedSection] = []
        full_text_parts: List[str] = []
        current_heading = "Introduction"

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            if para.style and para.style.name.startswith("Heading"):
                current_heading = text
                continue

            full_text_parts.append(text)
            sections.append(
                ParsedSection(
                    text=text,
                    heading=current_heading,
                    metadata={"heading": current_heading},
                )
            )

        # Parse tables as well
        for table in doc.tables:
            table_rows = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                table_rows.append(" | ".join(row_cells))
            if table_rows:
                table_text = "\n".join(table_rows)
                full_text_parts.append(table_text)
                sections.append(
                    ParsedSection(
                        text=table_text,
                        heading=f"{current_heading} (Table)",
                        metadata={"is_table": True, "heading": current_heading},
                    )
                )

        full_text = "\n\n".join(full_text_parts)
        return ParsedDocument(
            file_name=file_name,
            file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            full_text=full_text,
            sections=sections,
            page_count=1,
        )
